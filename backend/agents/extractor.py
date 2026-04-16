"""
Agente para extracción de texto de documentos
"""

from typing import Dict, Any
import io
from .base import BaseAgent
from pypdf import PdfReader
from docx import Document
import openpyxl
import logging

logger = logging.getLogger(__name__)


class DocumentExtractorAgent(BaseAgent):
    """
    Agente para extraer texto de documentos PDF, DOCX y Excel.
    
    No requiere LLM, usa librerías especializadas:
    - PyPDF para PDF
    - python-docx para Word
    - openpyxl para Excel
    """
    
    def __init__(self, **kwargs):
        super().__init__(llm=None, **kwargs)
    
    def validate_input(self, input_data: Dict[str, Any]) -> bool:
        """Valida que exista file_data y file_type"""
        required_keys = ["file_data", "file_type"]
        
        for key in required_keys:
            if key not in input_data:
                raise ValueError(f"Falta campo requerido: {key}")
        
        valid_types = ["pdf", "docx", "xlsx"]
        if input_data["file_type"] not in valid_types:
            raise ValueError(f"Tipo de archivo no soportado: {input_data['file_type']}")
        
        return True
    
    async def execute(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extrae texto del documento.
        
        Args:
            input_data: {
                "file_data": bytes,
                "file_type": str (pdf|docx|xlsx),
                "filename": str (opcional)
            }
            
        Returns:
            {
                "raw_text": str,
                "num_pages": int,
                "file_type": str,
                "metadata": dict
            }
        """
        self.validate_input(input_data)
        
        file_data = input_data["file_data"]
        file_type = input_data["file_type"]
        filename = input_data.get("filename", "unknown")
        
        self.log_execution("started", file_type=file_type, filename=filename)
        
        try:
            if file_type == "pdf":
                result = await self._extract_pdf(file_data)
            elif file_type == "docx":
                result = await self._extract_docx(file_data)
            elif file_type == "xlsx":
                result = await self._extract_excel(file_data)
            else:
                raise ValueError(f"Tipo no soportado: {file_type}")
            
            result["file_type"] = file_type
            result["filename"] = filename
            
            self.log_execution(
                "completed",
                chars_extracted=len(result["raw_text"]),
                num_pages=result.get("num_pages", 0)
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error extrayendo {file_type}: {str(e)}")
            raise
    
    async def _extract_pdf(self, file_data: bytes) -> Dict[str, Any]:
        """Extrae texto de PDF usando PyPDF"""
        pdf_file = io.BytesIO(file_data)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        return {
            "raw_text": "\n\n".join(text_parts),
            "num_pages": len(reader.pages),
            "metadata": {
                "pdf_version": reader.pdf_header,
                "encrypted": reader.is_encrypted
            }
        }
    
    async def _extract_docx(self, file_data: bytes) -> Dict[str, Any]:
        """Extrae texto de Word usando python-docx"""
        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extraer tablas también
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return {
            "raw_text": "\n".join(text_parts),
            "num_pages": len(doc.sections),
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
        }
    
    async def _extract_excel(self, file_data: bytes) -> Dict[str, Any]:
        """Extrae texto de Excel usando openpyxl"""
        excel_file = io.BytesIO(file_data)
        workbook = openpyxl.load_workbook(excel_file, data_only=True)
        
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n=== HOJA: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return {
            "raw_text": "\n".join(text_parts),
            "num_pages": len(workbook.sheetnames),
            "metadata": {
                "sheets": workbook.sheetnames,
                "num_sheets": len(workbook.sheetnames)
            }
        }
