"""Servicio de extracción de texto de documentos"""

import fitz  # PyMuPDF
from docx import Document
import openpyxl
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """Extrae texto de PDF, Word y Excel"""
    
    @staticmethod
    def extract(file_data: bytes, filename: str, content_type: str) -> Tuple[str, str]:
        """
        Extrae texto del documento
        
        Args:
            file_data: Bytes del archivo
            filename: Nombre del archivo
            content_type: MIME type
            
        Returns:
            Tuple(texto_extraido, tipo_documento)
        """
        try:
            if 'pdf' in content_type.lower():
                return DocumentExtractor._extract_pdf(file_data), "PDF"
            
            elif 'wordprocessingml' in content_type or 'msword' in content_type:
                return DocumentExtractor._extract_docx(file_data), "Word"
            
            elif 'spreadsheetml' in content_type or 'excel' in content_type:
                return DocumentExtractor._extract_excel(file_data), "Excel"
            
            else:
                raise ValueError(f"Tipo de archivo no soportado: {content_type}")
                
        except Exception as e:
            logger.error(f"Error extrayendo texto de {filename}: {str(e)}")
            raise
    
    @staticmethod
    def _extract_pdf(file_data: bytes) -> str:
        """Extrae texto de PDF usando PyMuPDF"""
        text_parts = []
        
        with fitz.open(stream=file_data, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Página {page_num} ---\n{text}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"✅ PDF extraído: {len(full_text)} caracteres, {len(text_parts)} páginas")
        return full_text
    
    @staticmethod
    def _extract_docx(file_data: bytes) -> str:
        """Extrae texto de Word (DOCX)"""
        from io import BytesIO
        
        doc = Document(BytesIO(file_data))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"✅ DOCX extraído: {len(full_text)} caracteres")
        return full_text
    
    @staticmethod
    def _extract_excel(file_data: bytes) -> str:
        """Extrae texto de Excel (XLSX)"""
        from io import BytesIO
        
        workbook = openpyxl.load_workbook(BytesIO(file_data), data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Hoja: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"✅ Excel extraído: {len(full_text)} caracteres, {len(workbook.sheetnames)} hojas")
        return full_text

# Instancia singleton
document_extractor = DocumentExtractor()
